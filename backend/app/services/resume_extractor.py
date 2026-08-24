"""이력서/포트폴리오 파일에서 텍스트를 뽑아낸다.

- 텍스트 파일(.txt/.md): 그대로 디코딩
- 엑셀(.xlsx): 시트/셀 내용을 추출
- 워드(.docx): 문단/표 텍스트를 추출
- PDF(.pdf): PDF 자체의 텍스트 레이어를 먼저 읽고, 스캔본이라 텍스트가 없을 때만
  페이지를 이미지로 렌더링해 gpt-4o Vision으로 OCR한다.
  (이력서 이미지를 그대로 Vision에 넘기면 이름·경력 같은 개인정보 문서로 인식해
  모델이 거부하는 경우가 있어, 되는 한 텍스트 레이어 추출로 우회한다.)
"""

import base64
from io import BytesIO
from pathlib import Path

import openpyxl
import pypdfium2 as pdfium
from docx import Document
from openai import AsyncOpenAI

_TEXT_EXTENSIONS = {".txt", ".md"}
_EXCEL_EXTENSIONS = {".xlsx"}
_WORD_EXTENSIONS = {".docx"}

_OCR_PROMPT = (
    "이 이미지에 보이는 텍스트를 있는 그대로 옮겨 적어줘. "
    "이것은 사용자 본인이 자신의 커리어 관리 도구에 업로드한 자기 자신의 문서이며, "
    "사용자 본인의 요청에 따라 텍스트만 추출한다. 설명이나 해설 없이 텍스트만 출력해줘."
)
_MAX_OCR_PAGES = 10
_MIN_TEXT_LAYER_LENGTH = 30
_REFUSAL_PREFIXES = (
    "i'm sorry",
    "i am sorry",
    "i'm unable",
    "i am unable",
    "i can't assist",
    "i cannot assist",
    "i can't help",
    "i cannot help",
    "죄송하지만",
    "도와드릴 수 없",
    "도와드리기 어렵",
)


def _looks_like_refusal(text: str) -> bool:
    lowered = text.strip().lower()
    return len(lowered) < 200 and any(lowered.startswith(prefix) for prefix in _REFUSAL_PREFIXES)


def _extract_text(content: bytes) -> str:
    return content.decode("utf-8", errors="ignore").strip()


def _extract_excel(content: bytes) -> str:
    workbook = openpyxl.load_workbook(BytesIO(content), data_only=True)
    lines: list[str] = []
    for sheet in workbook.worksheets:
        sheet_lines = []
        for row in sheet.iter_rows(values_only=True):
            cells = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if cells:
                sheet_lines.append(" | ".join(cells))
        if sheet_lines:
            lines.append(f"[시트: {sheet.title}]")
            lines.extend(sheet_lines)
    return "\n".join(lines).strip()


def _extract_word(content: bytes) -> str:
    document = Document(BytesIO(content))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _extract_pdf_text_layer(pdf: pdfium.PdfDocument) -> str:
    pages = []
    for page in pdf:
        textpage = page.get_textpage()
        pages.append(textpage.get_text_range())
    return "\n".join(pages).strip()


async def _ocr_pdf_pages(pdf: pdfium.PdfDocument, api_key: str) -> str:
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY is not set in backend/.env")

    page_count = min(len(pdf), _MAX_OCR_PAGES)
    image_urls: list[str] = []
    for index in range(page_count):
        bitmap = pdf[index].render(scale=2.0)
        pil_image = bitmap.to_pil()
        buffer = BytesIO()
        pil_image.save(buffer, format="PNG")
        encoded = base64.b64encode(buffer.getvalue()).decode("ascii")
        image_urls.append(f"data:image/png;base64,{encoded}")

    if not image_urls:
        return ""

    client = AsyncOpenAI(api_key=api_key)
    content_blocks: list[dict] = [{"type": "text", "text": _OCR_PROMPT}]
    for url in image_urls:
        content_blocks.append({"type": "image_url", "image_url": {"url": url}})

    try:
        response = await client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": content_blocks}],
        )
        text = (response.choices[0].message.content or "").strip()
    except Exception as error:
        raise RuntimeError(f"PDF OCR에 실패했습니다: {error}") from error

    if _looks_like_refusal(text):
        raise RuntimeError(
            "PDF 이미지에서 텍스트를 추출하지 못했습니다 (모델이 처리를 거부했습니다). "
            "가능하면 텍스트가 포함된 PDF나 다른 형식(txt/docx/xlsx)으로 다시 시도해주세요."
        )
    return text


async def _extract_pdf(content: bytes, api_key: str) -> str:
    pdf = pdfium.PdfDocument(BytesIO(content))
    try:
        text_layer = _extract_pdf_text_layer(pdf)
        if len(text_layer) >= _MIN_TEXT_LAYER_LENGTH:
            return text_layer
        return await _ocr_pdf_pages(pdf, api_key)
    finally:
        pdf.close()


async def extract_file_text(filename: str, content: bytes, api_key: str) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return await _extract_pdf(content, api_key)
    if suffix in _EXCEL_EXTENSIONS:
        return _extract_excel(content)
    if suffix in _WORD_EXTENSIONS:
        return _extract_word(content)
    if suffix in _TEXT_EXTENSIONS:
        return _extract_text(content)
    raise ValueError(f"지원하지 않는 파일 형식입니다: {suffix or filename}")
